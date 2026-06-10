"""
docx_utils.py — Word 파일 처리 유틸리티 모듈

python-docx를 사용한 Word 파일 생성 및 스타일링에 대한
재사용 가능한 유틸리티 함수를 제공합니다.
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def create_document() -> Any:
    """새로운 Word 문서를 생성합니다.
    
    Returns:
        Document 객체 (python-docx 라이브러리가 없는 경우 None)
    """
    if not _DOCX_AVAILABLE:
        return None
    return Document()


def add_heading(
    doc: Any,
    text: str,
    level: int = 1,
    alignment: str = "left"
) -> None:
    """문서에 제목을 추가합니다.
    
    Args:
        doc: Document 객체
        text: 제목 텍스트
        level: 제목 레벨 (1-3)
        alignment: 정렬 (left, center, right)
    """
    if not _DOCX_AVAILABLE or doc is None:
        return
    
    heading = doc.add_heading(text, level=level)
    
    if alignment == "center":
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(
    doc: Any,
    text: str,
    bold: bool = False,
    italic: bool = False,
    font_size: Optional[int] = None,
    color: Optional[str] = None
) -> None:
    """문서에 단락을 추가합니다.
    
    Args:
        doc: Document 객체
        text: 단락 텍스트
        bold: 굵게 여부
        italic: 기울임 여부
        font_size: 폰트 크기 (pt)
        color: 텍스트 색상 (RGB hex, 예: "FF0000")
    """
    if not _DOCX_AVAILABLE or doc is None:
        return
    
    paragraph = doc.add_paragraph(text)
    # 빈 단락인 경우 스타일 적용 건너뜀
    if not text or not paragraph.runs:
        return
    run = paragraph.runs[0]
    
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(
    doc: Any,
    data: List[List[Any]],
    headers: Optional[List[str]] = None,
    style: str = "Table Grid"
) -> None:
    """문서에 테이블을 추가합니다.
    
    Args:
        doc: Document 객체
        data: 테이블 데이터 (2D 리스트)
        headers: 헤더 행 (선택)
        style: 테이블 스타일
    """
    if not _DOCX_AVAILABLE or doc is None:
        return
    
    if headers:
        table_data = [headers] + data
    else:
        table_data = data
    
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = style
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)


def add_image(
    doc: Any,
    image_path: str | Path,
    width: Optional[float] = None,
    height: Optional[float] = None
) -> None:
    """문서에 이미지를 추가합니다.
    
    Args:
        doc: Document 객체
        image_path: 이미지 파일 경로
        width: 이미지 너비 (인치)
        height: 이미지 높이 (인치)
    """
    if not _DOCX_AVAILABLE or doc is None:
        return
    
    image_path = Path(image_path)
    if not image_path.exists():
        return
    
    kwargs = {}
    if width:
        kwargs["width"] = Inches(width)
    if height:
        kwargs["height"] = Inches(height)
    
    doc.add_picture(str(image_path), **kwargs)


def add_page_break(doc: Any) -> None:
    """문서에 페이지 나누기를 추가합니다.
    
    Args:
        doc: Document 객체
    """
    if not _DOCX_AVAILABLE or doc is None:
        return
    
    doc.add_page_break()


def save_document(doc: Any, file_path: str | Path) -> bool:
    """문서를 저장합니다.
    
    Args:
        doc: Document 객체
        file_path: 저장 경로
    
    Returns:
        저장 성공 여부
    """
    if not _DOCX_AVAILABLE or doc is None:
        return False
    
    try:
        doc.save(str(file_path))
        return True
    except Exception:
        return False


def generate_word_report(
    title: str,
    original_image_path: str | Path,
    restored_image_path: str | Path,
    measurements: Dict[str, Any],
    overall_score: float,
    perceived_age: Optional[int] = None,
    prescription: Optional[Dict[str, Any]] = None,
    llm_report: Optional[str] = None,
    llm_metric_opinions: Optional[List[Dict[str, Any]]] = None,
    llm_recommendations: Optional[List[str]] = None,
    llm_products: Optional[List[Dict[str, Any]]] = None,
    active_mixes: Optional[Dict[str, Any]] = None,
    ref_measurements: Optional[Dict[str, Any]] = None,
) -> Optional[Document]:
    """피부 분석 보고서를 Word 문서로 생성합니다.
    
    Args:
        title: 보고서 제목
        original_image_path: 원본 이미지 경로
        restored_image_path: 복원 이미지 경로
        measurements: 측정 결과 딕셔너리 (원본)
        overall_score: 종합 점수
        perceived_age: 인식 나이 (선택)
        prescription: 처방전 정보 (선택)
        llm_report: LLM 소견 (선택)
        llm_metric_opinions: LLM 측정 항목별 의견 (선택)
        llm_recommendations: LLM 추천 사항 (선택)
        llm_products: LLM 추천 제품 (선택)
        active_mixes: 활성 믹스 정보 (선택)
        ref_measurements: 측정 결과 딕셔너리 (기준/복원) (선택)
    
    Returns:
        Document 객체 (실패 시 None)
    """
    if not _DOCX_AVAILABLE:
        return None
    
    doc = create_document()
    if doc is None:
        return None
    
    # 제목
    add_heading(doc, title, level=1, alignment="center")
    
    # 타임스탬프
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y년 %m월 %d일 %H:%M:%S")
    add_paragraph(doc, f"생성일시: {timestamp}", font_size=10)
    add_paragraph(doc, "")  # 빈 줄
    
    # 이미지 섹션
    add_heading(doc, "이미지 비교", level=2)
    
    # 이미지 테이블 생성 (가로 배치)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.0)
    
    # 원본 이미지 셀
    cell_orig = table.cell(0, 0)
    cell_orig_paragraph = cell_orig.paragraphs[0]
    cell_orig_paragraph.add_run("원본 이미지").bold = True
    cell_orig_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if Path(original_image_path).exists():
        cell_orig_paragraph = cell_orig.add_paragraph()
        cell_orig_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_orig_paragraph.add_run()
        try:
            run.add_picture(str(original_image_path), width=Inches(3.0))
        except Exception:
            cell_orig.add_paragraph("이미지 로드 실패").italic = True
    else:
        cell_orig.add_paragraph("이미지를 찾을 수 없습니다.").italic = True
    
    # 복원 이미지 셀
    cell_restored = table.cell(0, 1)
    cell_restored_paragraph = cell_restored.paragraphs[0]
    cell_restored_paragraph.add_run("복원 이미지").bold = True
    cell_restored_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if Path(restored_image_path).exists():
        cell_restored_paragraph = cell_restored.add_paragraph()
        cell_restored_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_restored_paragraph.add_run()
        try:
            run.add_picture(str(restored_image_path), width=Inches(3.0))
        except Exception:
            cell_restored.add_paragraph("이미지 로드 실패").italic = True
    else:
        cell_restored.add_paragraph("이미지를 찾을 수 없습니다.").italic = True
    
    add_paragraph(doc, "")  # 빈 줄
    
    # 종합 점수 섹션
    add_heading(doc, "종합 점수", level=2)
    add_paragraph(doc, f"피부건강지수: {overall_score:.1f}", bold=True, font_size=14, color="0070C0")
    
    if perceived_age is not None:
        add_paragraph(doc, f"인식 나이: {perceived_age}세")
    
    add_paragraph(doc, "")  # 빈 줄
    
    # 점수등급 기준 섹션
    add_heading(doc, "점수등급 기준", level=2)
    
    grade_criteria = [
        ("90 이상", "매우 우수"),
        ("80~90", "우수"),
        ("70~80", "양호"),
        ("60~70", "집중케어 추천"),
        ("60 미만", "개선필요"),
    ]
    
    grade_table_data = [[range_str, grade] for range_str, grade in grade_criteria]
    add_table(doc, grade_table_data, headers=["점수 범위", "등급"])
    
    add_paragraph(doc, "")  # 빈 줄
    
    # 측정 결과 섹션
    add_heading(doc, "AI 측정 점수", level=2)
    
    table_data = []
    measurement_keys = []
    for key, value in measurements.items():
        measurement_keys.append(key)
        if isinstance(value, (int, float)):
            orig_score = f"{value:.1f}"
        else:
            orig_score = str(value)
        
        # 기준/복원 점수
        ref_score = ""
        if ref_measurements and key in ref_measurements:
            ref_value = ref_measurements[key]
            if isinstance(ref_value, (int, float)):
                ref_score = f"{ref_value:.1f}"
            else:
                ref_score = str(ref_value)
        
        table_data.append([key, orig_score, ref_score])
    
    if table_data:
        add_table(doc, table_data, headers=["항목명", "AI 측정 점수 (원본)", "AI 측정 점수 (기준)"])
    else:
        add_paragraph(doc, "측정 결과가 없습니다.", italic=True)
    
    # LLM 소견 섹션
    if llm_report:
        add_page_break(doc)
        add_heading(doc, "LLM 소견", level=2)
        add_paragraph(doc, llm_report)
    
    # LLM 측정 항목별 의견 섹션
    if llm_metric_opinions:
        add_page_break(doc)
        add_heading(doc, "측정 항목별 의견", level=2)
        
        # 측정 결과 테이블의 순서대로 정렬
        sorted_metrics = []
        for metric in llm_metric_opinions:
            # 객체인 경우 속성 접근, 딕셔너리인 경우 get 사용
            if hasattr(metric, 'display_name'):
                key = getattr(metric, 'key', '')
                display_name = getattr(metric, 'display_name', '')
            else:
                key = metric.get("key", "")
                display_name = metric.get("display_name", "")
            
            # 측정 결과 테이블의 키와 매칭
            if key in measurement_keys:
                sorted_metrics.append((measurement_keys.index(key), metric))
            elif display_name in measurement_keys:
                sorted_metrics.append((measurement_keys.index(display_name), metric))
            else:
                # 매칭되는 키가 없으면 마지막에 추가
                sorted_metrics.append((len(measurement_keys), metric))
        
        # 인덱스 순서대로 정렬
        sorted_metrics.sort(key=lambda x: x[0])
        
        for _, metric in sorted_metrics:
            # 객체인 경우 속성 접근, 딕셔너리인 경우 get 사용
            if hasattr(metric, 'display_name'):
                display_name = getattr(metric, 'display_name', '')
                score = getattr(metric, 'score', 0)
                opinion = getattr(metric, 'opinion', '')
                reason = getattr(metric, 'reason', '')
                grade = getattr(metric, 'grade', '')
            else:
                display_name = metric.get("display_name", metric.get("key", ""))
                score = metric.get("score", 0)
                opinion = metric.get("opinion", "")
                reason = metric.get("reason", "")
                grade = metric.get("grade", "")
            
            add_paragraph(doc, f"{display_name} ({score}점)", bold=True)
            if grade:
                add_paragraph(doc, f"등급: {grade}", font_size=9)
            if opinion:
                add_paragraph(doc, f"의견: {opinion}")
            if reason:
                add_paragraph(doc, f"이유: {reason}")
            add_paragraph(doc, "")  # 빈 줄
    
    # LLM 추천 사항 섹션
    if llm_recommendations:
        add_page_break(doc)
        add_heading(doc, "추천 사항", level=2)
        
        for i, recommendation in enumerate(llm_recommendations, 1):
            add_paragraph(doc, f"{i}. {recommendation}")
    
    # LLM 추천 제품 섹션
    if llm_products:
        add_page_break(doc)
        add_heading(doc, "추천 제품", level=2)
        
        for product in llm_products:
            # 객체인 경우 속성 접근, 딕셔너리인 경우 get 사용
            if hasattr(product, 'product_name'):
                product_name = getattr(product, 'product_name', '')
                category = getattr(product, 'category', '')
                efficacy = getattr(product, 'efficacy', '')
                key_ingredients = getattr(product, 'key_ingredients', [])
            else:
                product_name = product.get("product_name", "")
                category = product.get("category", "")
                efficacy = product.get("efficacy", "")
                key_ingredients = product.get("key_ingredients", [])
            
            add_paragraph(doc, f"제품명: {product_name}", bold=True)
            add_paragraph(doc, f"카테고리: {category}")
            add_paragraph(doc, f"효능: {efficacy}")
            
            if key_ingredients:
                add_paragraph(doc, f"주요 성분: {', '.join(key_ingredients)}")
            
            add_paragraph(doc, "")  # 빈 줄
    
    # 처방전 섹션 (제일 마지막)
    add_page_break(doc)
    add_heading(doc, "처방전 (Prescription)", level=2)
    # 처방전 섹션은 비워둠 (상세 테이블에 표시되므로)
    
    # 활성 믹스 테이블 (M01-M13)
    add_heading(doc, "활성 믹스 (M01-M13)", level=2)
    
    # 테이블 데이터 생성 (M01-M13 모두 표시)
    table_data = []
    if active_mixes:
        for mix_code in sorted(active_mixes.keys()):
            if mix_code.startswith("M") and mix_code != "_note":
                mix_info = active_mixes[mix_code]
                name = mix_info.get("name", "")
                category = mix_info.get("category", "")
                description = mix_info.get("description", "")
                ingredients = ", ".join(mix_info.get("ingredients", []))
                
                # 처방전에서 배합비 추출
                percentage = ""
                if isinstance(prescription, dict):
                    # assessment에서 배합비 추출
                    if "assessment" in prescription:
                        assessment = prescription.get("assessment", {})
                        if mix_code in assessment:
                            mix_data = assessment[mix_code]
                            if isinstance(mix_data, dict):
                                percentage = f"{mix_data.get('percentage', 0)}%"
                            else:
                                percentage = f"{mix_data}%"
                    # M01(베이스믹스)인 경우 base 섹션에서도 배합비 추출
                    if mix_code == "M01" and "base" in prescription:
                        base_data = prescription.get("base", {})
                        if isinstance(base_data, dict):
                            percentage = f"{base_data.get('percentage', 0)}%"
                        else:
                            percentage = f"{base_data}%"
                
                table_data.append([
                    mix_code,
                    name,
                    category,
                    description,
                    ingredients,
                    percentage
                ])
    
    if table_data:
        add_table(doc, table_data, headers=["코드", "이름", "카테고리", "설명", "주요 성분", "배합비"])
    else:
        add_paragraph(doc, "활성 믹스 정보가 없습니다.", italic=True)
    
    return doc
